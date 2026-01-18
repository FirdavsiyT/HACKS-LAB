import csv
import codecs
import uuid
from django import forms
from django.core.cache import cache
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.views.decorators.http import require_POST
from django.db.models import Count, Sum
from django.db.models.functions import Coalesce
from django.db import transaction
from django.core.exceptions import PermissionDenied
from django.utils import timezone
from datetime import timedelta
from .models import LessonTemplate
from .forms import  LessonTemplateForm
# Try to import python-docx, handle gracefully if missing
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    Document = None

from pages.models import Challenge, Solve, Category, Attempt
from users.models import User
from .models import LessonSettings
from .forms import ChallengeForm, CategoryForm, TimerSettingsForm


# --- Custom Decorator ---
def mentor_required(view_func):
    def _wrapped_view(request, *args, **kwargs):
        if request.user.is_superuser or request.user.groups.filter(name='Mentors').exists():
            return view_func(request, *args, **kwargs)
        else:
            raise PermissionDenied

    return _wrapped_view


# --- FORMS ---

class MentorMessageForm(forms.Form):
    # Field is validated but rendered manually in template for Avatars
    recipient = forms.CharField(widget=forms.HiddenInput())
    message = forms.CharField(widget=forms.Textarea(attrs={
        'rows': 4,
        'class': 'w-full bg-[#1a1c23] border border-[#2c2f3b] rounded p-3 text-white focus:border-[#9fef00] focus:outline-none transition-colors'
    }), label="Message Content")


# --- VIEWS ---

@login_required
@mentor_required
def dashboard(request):
    total_users = User.objects.filter(is_superuser=False).exclude(groups__name='Mentors').count()
    total_challenges = Challenge.objects.count()
    total_solves = Solve.objects.count()

    recent_solves = Solve.objects.select_related('user', 'challenge').order_by('-date')[:10]

    # Timer logic
    lesson_settings = LessonSettings.get_settings()

    initial_data = {}
    if lesson_settings.start_time and lesson_settings.end_time:
        duration = (lesson_settings.end_time - lesson_settings.start_time).total_seconds() / 60
        delay = 0
        if lesson_settings.hard_deadline:
            delay = (lesson_settings.hard_deadline - lesson_settings.end_time).total_seconds() / 60

        initial_data = {
            'duration_minutes': int(round(duration)),
            'delay_minutes': int(round(delay))
        }

    timer_form = TimerSettingsForm(initial=initial_data)

    if request.method == 'POST' and 'set_timer' in request.POST:
        timer_form = TimerSettingsForm(request.POST)
        if timer_form.is_valid():
            minutes = timer_form.cleaned_data.get('duration_minutes')
            delay = timer_form.cleaned_data.get('delay_minutes') or 0

            if minutes:
                now = timezone.now()

                if lesson_settings.start_time and lesson_settings.end_time:
                    lesson_settings.end_time = lesson_settings.start_time + timedelta(minutes=minutes)
                    action_msg = f"Timer updated. Total duration: {minutes} min."
                else:
                    lesson_settings.start_time = now
                    lesson_settings.end_time = now + timedelta(minutes=minutes)
                    action_msg = f"Timer started for {minutes} minutes."

                lesson_settings.hard_deadline = lesson_settings.end_time + timedelta(minutes=delay)
                lesson_settings.save()
                messages.success(request, f'{action_msg} Hard stop delay: {delay} min.')
            return redirect('mentors:dashboard')

    if request.method == 'POST' and 'reset_timer' in request.POST:
        lesson_settings.start_time = None
        lesson_settings.end_time = None
        lesson_settings.hard_deadline = None
        lesson_settings.save()
        messages.success(request, 'Timer stopped.')
        return redirect('mentors:dashboard')

    context = {
        'total_users': total_users,
        'total_challenges': total_challenges,
        'total_solves': total_solves,
        'recent_solves': recent_solves,
        'lesson_settings': lesson_settings,
        'timer_form': timer_form,
    }
    return render(request, 'mentors/dashboard.html', context)


# --- MESSAGING SYSTEM (With Avatar Support) ---

@login_required
@mentor_required
def send_message(request):
    """
    View for mentors to send ephemeral messages via Cache.
    """
    # Fetch ALL active users (Students, Mentors, Admins)
    users_list = User.objects.filter(is_active=True).order_by('username')

    if request.method == 'POST':
        form = MentorMessageForm(request.POST)
        if form.is_valid():
            recipient = form.cleaned_data['recipient']
            content = form.cleaned_data['message']
            msg_id = str(uuid.uuid4())

            message_data = {
                'id': msg_id,
                'text': content,
                'timestamp': timezone.now().timestamp(),
                'sender': request.user.username
            }

            if recipient == 'all':
                broadcasts = cache.get('hacklabs_broadcasts', [])
                broadcasts.append(message_data)
                if len(broadcasts) > 20:
                    broadcasts = broadcasts[-20:]
                cache.set('hacklabs_broadcasts', broadcasts, timeout=120)
                messages.success(request, "Broadcast sent to ALL users!")
            else:
                user_key = f'hacklabs_msgs_{recipient}'
                user_msgs = cache.get(user_key, [])
                user_msgs.append(message_data)
                cache.set(user_key, user_msgs, timeout=300)
                messages.success(request, "Message sent to user.")

            return redirect('mentors:send_message')
    else:
        # Set default to 'all'
        form = MentorMessageForm(initial={'recipient': 'all'})

    return render(request, 'mentors/message_form.html', {
        'form': form,
        'title': 'Send Message',
        'users_list': users_list  # Renamed from 'students' to 'users_list'
    })


@login_required
def check_messages(request):
    if not request.user.is_authenticated:
        return JsonResponse({'messages': []})

    new_messages = []

    user_key = f'hacklabs_msgs_{request.user.id}'
    personal_msgs = cache.get(user_key, [])

    if personal_msgs:
        for msg in personal_msgs:
            msg['type'] = 'personal'
            new_messages.append(msg)
        cache.delete(user_key)

    broadcasts = cache.get('hacklabs_broadcasts', [])
    for msg in broadcasts:
        msg['type'] = 'broadcast'
        new_messages.append(msg)

    return JsonResponse({'messages': new_messages})


# --- CHALLENGES ---

@login_required
@mentor_required
def challenges_list(request):
    sort_by = request.GET.get('sort', 'newest')
    category_filter = request.GET.get('category')

    challenges = Challenge.objects.select_related('category')

    if category_filter:
        challenges = challenges.filter(category__name=category_filter)

    if sort_by == 'category':
        challenges = challenges.order_by('category__name', 'title')
    elif sort_by == 'points_asc':
        challenges = challenges.order_by('points')
    elif sort_by == 'points_desc':
        challenges = challenges.order_by('-points')
    elif sort_by == 'active':
        challenges = challenges.order_by('-is_active', '-id')
    elif sort_by == 'inactive':
        challenges = challenges.order_by('is_active', '-id')
    else:
        challenges = challenges.order_by('-id')

    categories = Category.objects.all().order_by('name')

    context = {
        'challenges': challenges,
        'categories': categories,
        'current_sort': sort_by,
        'current_category': category_filter
    }
    return render(request, 'mentors/challenges_list.html', context)


@login_required
@mentor_required
def export_challenges_docx(request):
    if Document is None:
        messages.error(request,
                       "Library 'python-docx' is missing. Please install it on the server: pip install python-docx")
        return redirect('mentors:challenges_list')

    challenges = Challenge.objects.filter(is_active=True).select_related('category').order_by('category__name',
                                                                                              'points')

    if not challenges.exists():
        messages.warning(request, "No active challenges to export.")
        return redirect('mentors:challenges_list')

    document = Document()
    main_title = document.add_heading('HackLabs Active Challenges', 0)
    main_title.alignment = 1

    timestamp = timezone.now().strftime('%Y-%m-%d %H:%M')
    p_date = document.add_paragraph(f"Exported on: {timestamp}")
    p_date.alignment = 1
    document.add_paragraph("-" * 50).alignment = 1

    for challenge in challenges:
        p = document.add_paragraph()
        p.space_after = Pt(12)

        def add_line(label, text):
            run_label = p.add_run(f"{label} ")
            run_label.bold = True
            run_label.font.color.rgb = RGBColor(0, 0, 0)
            p.add_run(f"{text}\n")

        add_line("Title:", challenge.title)
        add_line("Category:", challenge.category.name if challenge.category else "Uncategorized")
        add_line("Description:", challenge.description)
        add_line("Flag:", challenge.flag)
        add_line("Award:", f"{challenge.points} Pts")

        sep = document.add_paragraph("_" * 30)
        sep.alignment = 1
        sep.space_after = Pt(24)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"challenges_export_{timezone.now().strftime('%Y%m%d')}.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    document.save(response)
    return response


@login_required
@mentor_required
@require_POST
def bulk_challenges_action(request):
    action = request.POST.get('action')
    challenge_ids = request.POST.getlist('challenge_ids')

    if not challenge_ids:
        messages.warning(request, 'No challenges selected.')
        return redirect('mentors:challenges_list')

    if action == 'enable_selected':
        count = Challenge.objects.filter(id__in=challenge_ids).update(is_active=True)
        messages.success(request, f'Enabled {count} challenges.')
    elif action == 'disable_selected':
        count = Challenge.objects.filter(id__in=challenge_ids).update(is_active=False)
        messages.success(request, f'Disabled {count} challenges.')

    return redirect('mentors:challenges_list')


@login_required
@mentor_required
def challenge_create(request):
    if request.method == 'POST':
        form = ChallengeForm(request.POST)
        if form.is_valid():
            challenge = form.save(commit=False)
            challenge.author = request.user.username
            challenge.save()
            messages.success(request, 'Challenge created successfully!')
            return redirect('mentors:challenges_list')
    else:
        form = ChallengeForm()

    return render(request, 'mentors/challenge_form.html', {'form': form, 'title': 'Create Challenge'})


@login_required
@mentor_required
def challenge_edit(request, pk):
    challenge = get_object_or_404(Challenge, pk=pk)
    if request.method == 'POST':
        form = ChallengeForm(request.POST, instance=challenge)
        if form.is_valid():
            form.save()
            messages.success(request, 'Challenge updated successfully!')
            return redirect('mentors:challenges_list')
    else:
        form = ChallengeForm(instance=challenge)

    return render(request, 'mentors/challenge_form.html', {'form': form, 'title': 'Edit Challenge'})


@login_required
@mentor_required
def challenge_delete(request, pk):
    challenge = get_object_or_404(Challenge, pk=pk)
    if request.method == 'POST':
        challenge.delete()
        messages.success(request, 'Challenge deleted!')
        return redirect('mentors:challenges_list')
    return render(request, 'mentors/challenge_confirm_delete.html', {'object': challenge, 'type': 'Challenge'})


@login_required
@mentor_required
@require_POST
def challenge_toggle_active(request, pk):
    challenge = get_object_or_404(Challenge, pk=pk)
    challenge.is_active = not challenge.is_active
    challenge.save()
    messages.success(request, f'Challenge "{challenge.title}" is now {"Active" if challenge.is_active else "Hidden"}')
    return redirect('mentors:challenges_list')


@login_required
@mentor_required
@require_POST
def disable_all_challenges(request):
    if request.POST.get('confirm') != 'CONFIRM_DISABLE':
        messages.error(request, 'Confirmation failed. Type CONFIRM_DISABLE exactly.')
        return redirect('mentors:challenges_list')

    updated_count = Challenge.objects.update(is_active=False)

    messages.success(request, f'Lockdown initiated! {updated_count} challenges hidden.')
    return redirect('mentors:challenges_list')


# --- CATEGORIES ---

@login_required
@mentor_required
def categories_list(request):
    categories = Category.objects.annotate(challenge_count=Count('challenges')).order_by('name')
    return render(request, 'mentors/categories_list.html', {'categories': categories})


@login_required
@mentor_required
def category_create(request):
    if request.method == 'POST':
        form = CategoryForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Category created!')
            return redirect('mentors:categories_list')
    else:
        form = CategoryForm()
    return render(request, 'mentors/category_form.html', {'form': form, 'title': 'Create Category'})


@login_required
@mentor_required
def category_edit(request, pk):
    category = get_object_or_404(Category, pk=pk)
    if request.method == 'POST':
        form = CategoryForm(request.POST, instance=category)
        if form.is_valid():
            form.save()
            messages.success(request, 'Category updated!')
            return redirect('mentors:categories_list')
    else:
        form = CategoryForm(instance=category)
    return render(request, 'mentors/category_form.html', {'form': form, 'title': 'Edit Category'})


@login_required
@mentor_required
def category_delete(request, pk):
    category = get_object_or_404(Category, pk=pk)
    if request.method == 'POST':
        category.delete()
        messages.success(request, 'Category deleted!')
        return redirect('mentors:categories_list')
    return render(request, 'mentors/category_confirm_delete.html', {'object': category, 'type': 'Category'})


# --- USERS & SYSTEM ---

@login_required
@mentor_required
def users_list(request):
    max_points_data = Challenge.objects.filter(is_active=True).aggregate(total=Sum('points'))
    max_possible_points = max_points_data['total'] if max_points_data['total'] is not None else 0

    users_qs = User.objects.filter(is_superuser=False).exclude(groups__name='Mentors').annotate(
        total_points=Coalesce(Sum('solves__challenge__points'), 0),
        solved_count=Count('solves')
    ).order_by('-total_points')

    users = []
    for user in users_qs:
        if max_possible_points > 0:
            user.percentage = (user.total_points / max_possible_points) * 100
        else:
            user.percentage = 0.0
        users.append(user)

    context = {
        'users': users,
        'max_possible_points': max_possible_points
    }
    return render(request, 'mentors/users_list.html', context)


@login_required
@mentor_required
def export_users_csv(request):
    response = HttpResponse(content_type='text/csv')
    timestamp = timezone.now().strftime('%Y-%m-%d_%H-%M')
    response['Content-Disposition'] = f'attachment; filename="hacklabs_results_{timestamp}.csv"'

    response.write(codecs.BOM_UTF8)
    writer = csv.writer(response, delimiter=';', quotechar='"', quoting=csv.QUOTE_MINIMAL)
    writer.writerow(['Rank', 'Username', 'Total Score', 'Max Possible Score', 'Completion Percentage'])

    max_points_data = Challenge.objects.filter(is_active=True).aggregate(total=Sum('points'))
    max_possible_points = max_points_data['total'] if max_points_data['total'] is not None else 0

    users = User.objects.filter(is_superuser=False).exclude(groups__name='Mentors').annotate(
        total_points=Coalesce(Sum('solves__challenge__points'), 0)
    ).order_by('-total_points')

    for index, user in enumerate(users, 1):
        user_points = user.total_points
        if max_possible_points > 0:
            percentage = (user_points / max_possible_points) * 100
        else:
            percentage = 0.0
        percentage_str = f"{percentage:.2f}%"
        writer.writerow([index, user.username, user_points, max_possible_points, percentage_str])

    return response


@login_required
@mentor_required
@require_POST
def reset_platform(request):
    if request.POST.get('confirm') != 'CONFIRM_RESET':
        messages.error(request, 'Confirmation failed. Type CONFIRM_RESET exactly.')
        return redirect('mentors:users_list')

    with transaction.atomic():
        Solve.objects.all().delete()
        Attempt.objects.all().delete()

        lesson_settings = LessonSettings.get_settings()
        lesson_settings.start_time = None
        lesson_settings.end_time = None
        lesson_settings.hard_deadline = None
        lesson_settings.save()

        deleted_count, _ = User.objects.filter(
            is_superuser=False,
            is_staff=False
        ).exclude(
            groups__name='Mentors'
        ).delete()

        messages.success(request, f'Platform Reset Complete! Removed {deleted_count} users and all progress.')

    return redirect('mentors:users_list')


# --- LESSON TEMPLATES (NEW) ---

@login_required
@mentor_required
def templates_list(request):
    """
    Shows all created lesson templates.
    """
    templates = LessonTemplate.objects.annotate(
        task_count=Count('challenges'),
        total_score=Coalesce(Sum('challenges__points'), 0)
    ).order_by('-created_at')

    return render(request, 'mentors/lesson_templates_list.html', {'templates': templates})


@login_required
@mentor_required
def template_create(request):
    if request.method == 'POST':
        form = LessonTemplateForm(request.POST)
        if form.is_valid():
            template = form.save()
            messages.success(request, f'Template "{template.title}" created!')
            return redirect('mentors:templates_list')
    else:
        form = LessonTemplateForm()

    # Categorize challenges for the view to render grouped checkboxes
    challenges_by_category = {}
    challenges = Challenge.objects.all().select_related('category').order_by('category__name', 'points')
    for ch in challenges:
        cat_name = ch.category.name if ch.category else 'Uncategorized'
        if cat_name not in challenges_by_category:
            challenges_by_category[cat_name] = []
        challenges_by_category[cat_name].append(ch)

    return render(request, 'mentors/lesson_template_form.html', {
        'form': form,
        'challenges_by_category': challenges_by_category,
        'title': 'Create Lesson Template'
    })


@login_required
@mentor_required
def template_edit(request, pk):
    template = get_object_or_404(LessonTemplate, pk=pk)
    if request.method == 'POST':
        form = LessonTemplateForm(request.POST, instance=template)
        if form.is_valid():
            form.save()
            messages.success(request, f'Template "{template.title}" updated!')
            return redirect('mentors:templates_list')
    else:
        form = LessonTemplateForm(instance=template)

    # Categorize challenges for display
    challenges_by_category = {}
    challenges = Challenge.objects.all().select_related('category').order_by('category__name', 'points')
    for ch in challenges:
        cat_name = ch.category.name if ch.category else 'Uncategorized'
        if cat_name not in challenges_by_category:
            challenges_by_category[cat_name] = []
        challenges_by_category[cat_name].append(ch)

    return render(request, 'mentors/lesson_template_form.html', {
        'form': form,
        'challenges_by_category': challenges_by_category,
        'title': 'Edit Lesson Template',
        'template': template
    })


@login_required
@mentor_required
@require_POST
def template_delete(request, pk):
    template = get_object_or_404(LessonTemplate, pk=pk)
    template.delete()
    messages.success(request, 'Template deleted.')
    return redirect('mentors:templates_list')


@login_required
@mentor_required
@require_POST
def template_apply(request, pk):
    """
    Core Logic: Applies the challenges from the template to the live platform.
    Action types:
    - 'exclusive': Disable ALL other challenges, enable ONLY template challenges.
    - 'enable': Enable template challenges (others remain unchanged).
    - 'disable': Disable template challenges.
    """
    template = get_object_or_404(LessonTemplate, pk=pk)
    action = request.POST.get('action')

    with transaction.atomic():
        template_challenges = template.challenges.all()

        if action == 'exclusive':
            # Hide EVERYTHING first
            Challenge.objects.update(is_active=False)
            # Enable template items
            count = template_challenges.update(is_active=True)
            messages.success(request, f'LESSON STARTED: "{template.title}". Enabled {count} tasks. All others hidden.')

        elif action == 'enable':
            count = template_challenges.update(is_active=True)
            messages.success(request, f'Enabled {count} tasks from "{template.title}".')

        elif action == 'disable':
            count = template_challenges.update(is_active=False)
            messages.success(request, f'Disabled {count} tasks from "{template.title}".')

    return redirect('mentors:templates_list')